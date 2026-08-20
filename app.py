# -*- coding: utf-8 -*-
"""
TTS Studio local (mô phỏng giao diện Vbee Studio TTS)
Chạy:  python app.py  -> mở http://127.0.0.1:5000
Dùng pipeline TTS giọng Ngọc Huyền có sẵn (piper_ngochuyen).
"""
import os
import re
import io
import sys
import uuid
import json
import shutil
import tempfile
import threading
import subprocess
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, jsonify, send_from_directory

# ---- Đường dẫn pipeline TTS (dùng lại module có sẵn) ----
PIPELINE_DIR = Path(r"C:\Users\Khang\Desktop\dịch tryện\Nghiên cứu giọng kênh\piper_ngochuyen")
if str(PIPELINE_DIR) not in sys.path:
    sys.path.insert(0, str(PIPELINE_DIR))

import tts_ngochuyen as tts  # noqa: E402
import loanwords_dynamic  # noqa: E402 — kho từ điển động tự học
from piper.config import SynthesisConfig  # noqa: E402

BASE_DIR = Path(__file__).parent
PRODUCT_DIR = BASE_DIR / "sản phẩm"
PRODUCT_DIR.mkdir(exist_ok=True)
VIDEO_DIR = BASE_DIR / "video_output"
VIDEO_DIR.mkdir(exist_ok=True)
UPLOAD_IMG_DIR = BASE_DIR / "upload_images"
UPLOAD_IMG_DIR.mkdir(exist_ok=True)

# ---- FFprobe để lấy duration audio ----
FFPROBE = r"C:\Users\Khang\Desktop\ffmpeg\ffmpeg-8.1.2-essentials_build\bin\ffprobe.exe"


def get_audio_duration(filepath: Path) -> float:
    """Lấy duration (giây) của file audio bằng ffprobe."""
    try:
        r = subprocess.run(
            [FFPROBE, "-v", "error", "-show_entries", "format=duration",
             "-of", "default=noprint_wrappers=1:nokey=1", str(filepath)],
            capture_output=True, text=True, timeout=5,
        )
        return float(r.stdout.strip()) if r.stdout.strip() else 0.0
    except Exception:
        return 0.0


# ---- Bản nháp kết quả quét từ nước ngoài ----
# Lưu ở server để F5 / tải lại trang KHÔNG làm mất các từ đã quét và cách đọc
# LO vừa sửa tay. Chỉ xóa khi LO bấm "Xóa kết quả quét" hoặc thêm hết vào kho.
SCAN_DRAFT_FILE = BASE_DIR / "scan_draft.json"
SCAN_DRAFT_LOCK = threading.Lock()


def _load_scan_draft() -> list:
    try:
        if SCAN_DRAFT_FILE.exists():
            return json.loads(SCAN_DRAFT_FILE.read_text(encoding="utf-8")) or []
    except Exception:
        pass
    return []


def _save_scan_draft(items: list):
    with SCAN_DRAFT_LOCK:
        SCAN_DRAFT_FILE.write_text(
            json.dumps(items, ensure_ascii=False, indent=1),
            encoding="utf-8",
        )


def _clear_scan_draft():
    with SCAN_DRAFT_LOCK:
        if SCAN_DRAFT_FILE.exists():
            SCAN_DRAFT_FILE.unlink()

app = Flask(__name__)

# ---- Voice cache: load model 1 lần, render nhiều lần ----
_voice = None
_voice_lock = threading.Lock()

def get_voice():
    global _voice
    with _voice_lock:
        if _voice is None:
            from piper.voice import PiperVoice
            v = PiperVoice.load(str(tts.MODEL))
            tts.patch_anh_phoneme(v)
            tts.patch_phatam_fix(v)
            _voice = v
        return _voice


# ---- Quản lý tiến trình render ----
RENDERS = {}   # id -> {status, progress, total, message, error, file, text_preview}
RENDER_LOCK = threading.Lock()
_rendering = False
_render_thread = None


def _eq_filter(eq: str):
    if eq == "strong":
        return "highpass=f=80,equalizer=f=1800:t=q:w=1.2:g=4,equalizer=f=4500:t=q:w=1.5:g=3.5,equalizer=f=9000:t=q:w=1.5:g=2,lowpass=f=11000"
    if eq == "accent":
        return "highpass=f=60,equalizer=f=300:t=q:w=1.2:g=3.5,equalizer=f=1200:t=q:w=1.2:g=3,equalizer=f=2600:t=q:w=1.5:g=2.5,equalizer=f=5000:t=q:w=1.5:g=1.5,lowpass=f=12000"
    if eq == "vbee":
        return ("highpass=f=120,"
                "equalizer=f=800:t=q:w=1.5:g=2,"
                "equalizer=f=2500:t=q:w=1.2:g=3.5,"
                "equalizer=f=6000:t=q:w=1.8:g=2,"
                "lowpass=f=12000,"
                "loudnorm=I=-14:TP=-1.5:LRA=9")
    return "loudnorm=I=-16:TP=-1.5:LRA=11"


def build_filter_chain(eq: str, volume: float, pitch_semitones: float) -> str:
    """Chuỗi filter ffmpeg: đổi cao độ giọng (pitch) giữ nguyên tốc độ + EQ + âm lượng.

    pitch: số bán âm (-12..+12). Dùng asetrate nhân mẫu theo 2^(n/12) rồi atempo bù
    ngược lại để giữ nguyên tốc độ đọc; aresample về 22050 chuẩn.
    volume: hệ số âm lượng (0.5..2.0).
    """
    factor = 2.0 ** (pitch_semitones / 12.0)
    chains = []
    if abs(pitch_semitones) > 0.001:
        chains.append(f"asetrate={int(tts.SAMPLE_RATE * factor)},atempo={1.0 / factor},aresample={tts.SAMPLE_RATE}")
    chains.append(_eq_filter(eq))
    if abs(volume - 1.0) > 0.001:
        chains.append(f"volume={volume:.3f}")
    return ",".join(chains)


def render_job(rid: str, text: str, settings: dict):
    global _rendering
    try:
        RENDERS[rid]["status"] = "processing"
        RENDERS[rid]["message"] = "Đang xử lý văn bản..."
        # Xóa chú thích #... trước khi xử lý
        text = re.sub(r"#[^\n]*", "", text).strip()
        # Render với kho đọc hiện tại — tuyệt đối không tự ý thêm từ mới vào kho
        text_proc = tts.vietnamize_text(text)
        sentences = tts.split_into_sentences(text_proc)
        total = len(sentences)
        RENDERS[rid]["total"] = total

        speed = float(settings.get("speed", 1.3))
        pause = float(settings.get("pause", 0.18))
        pause_comma = float(settings.get("pause_comma", 0.18))
        noise_scale = float(settings.get("noise_scale", 0.667))
        noise_w = float(settings.get("noise_w", 0.8))
        eq = settings.get("eq", "none")
        volume = float(settings.get("volume", 1.0))
        pitch = float(settings.get("pitch", 0.0))

        syn = SynthesisConfig(
            length_scale=1.0 / speed,
            noise_scale=noise_scale,
            noise_w_scale=noise_w,
        )
        voice = get_voice()

        parts = []
        done = 0
        for i, s in enumerate(sentences):
            if not re.search(r"[A-Za-zÀ-ỹ0-9]", s):
                continue
            try:
                wav = tts.synth_sentence(voice, syn, s)
            except Exception:
                continue
            parts.append(wav)
            last_ch = s[-1] if s else ""
            if last_ch in "，,。.":
                parts.append(tts.silence(pause_comma if last_ch in "，," else pause))
            done += 1
            RENDERS[rid]["progress"] = done
            RENDERS[rid]["message"] = f"Đang đọc câu {done}/{total}..."
            if done % 5 == 0:
                RENDERS[rid]["message"] = f"Đang đọc câu {done}/{total}..."

        if not parts:
            raise RuntimeError("Không có câu nào đọc được (text trống?)")

        full = b"".join(parts)
        RENDERS[rid]["message"] = "Đang ghi audio..."

        out_name = f"{datetime.now():%Y%m%d_%H%M%S}_{rid[:8]}.mp3"
        out_path = PRODUCT_DIR / out_name

        tmp = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
        tmp.close()
        tts.write_wav_pcm(full, tmp.name)

        ff = tts.FFMPEG
        if not ff:
            raise RuntimeError("Không tìm thấy ffmpeg để xuất MP3")
        ff_args = [ff, "-y", "-i", tmp.name, "-codec:a", "libmp3lame", "-qscale:a", "2",
                   "-af", build_filter_chain(eq, volume, pitch)]
        subprocess.run(ff_args + [str(out_path)], capture_output=True)
        os.unlink(tmp.name)

        RENDERS[rid]["status"] = "done"
        RENDERS[rid]["message"] = "Hoàn tất"
        RENDERS[rid]["file"] = out_name
    except Exception as e:
        RENDERS[rid]["status"] = "error"
        RENDERS[rid]["error"] = str(e)
    finally:
        with RENDER_LOCK:
            _rendering = False


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/tts", methods=["POST"])
def api_tts():
    global _rendering, _render_thread
    data = request.get_json(force=True, silent=True) or {}
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "Vui lòng nhập văn bản"}), 400

    with RENDER_LOCK:
        if _rendering:
            return jsonify({"error": "Đang render audio khác, chờ xong rồi thử lại"}), 409
        _rendering = True

    rid = uuid.uuid4().hex
    RENDERS[rid] = {
        "status": "queued", "progress": 0, "total": 0,
        "message": "Bắt đầu...", "error": None, "file": None,
        "text_preview": text[:80],
    }
    settings = data.get("settings") or {}
    _render_thread = threading.Thread(target=render_job, args=(rid, text, settings), daemon=True)
    _render_thread.start()
    return jsonify({"id": rid})


@app.route("/api/status/<rid>")
def api_status(rid):
    r = RENDERS.get(rid)
    if not r:
        return jsonify({"error": "Không tìm thấy job"}), 404
    return jsonify({
        "status": r["status"],
        "progress": r["progress"],
        "total": r["total"],
        "message": r["message"],
        "error": r["error"],
        "file": r["file"],
    })


@app.route("/api/loanwords")
def api_loanwords():
    """Kho từ điển động: số từ + danh sách + cách đọc (từ thêm sau nằm ở TRÊN CÙNG)."""
    data = loanwords_dynamic.get_all()
    items = list(reversed([{"word": k, "reading": v} for k, v in data.items()]))
    return jsonify({"total": len(items), "items": items})


@app.route("/api/loanwords", methods=["POST"])
def api_loanwords_add():
    """Thêm/sửa nhiều từ vào kho động. body: {items: [{word, reading}]}."""
    data = request.get_json(force=True, silent=True) or {}
    items = data.get("items") or []
    pairs = [(it.get("word", ""), it.get("reading", "")) for it in items]
    added = loanwords_dynamic.add_words(pairs)
    return jsonify({"ok": True, "added": added, "total": loanwords_dynamic.count()})


@app.route("/api/loanwords/<word>", methods=["PUT"])
def api_loanwords_update(word):
    """Sửa cách đọc một từ. body: {reading: "..."}."""
    data = request.get_json(force=True, silent=True) or {}
    reading = data.get("reading", "")
    ok = loanwords_dynamic.update_word(word, reading)
    if not ok:
        return jsonify({"error": "Không tìm thấy từ"}), 404
    return jsonify({"ok": True})


@app.route("/api/loanwords/<word>", methods=["DELETE"])
def api_loanwords_delete(word):
    ok = loanwords_dynamic.remove_word(word)
    if not ok:
        return jsonify({"error": "Không tìm thấy từ"}), 404
    return jsonify({"ok": True})


@app.route("/api/scan", methods=["POST"])
def api_scan():
    """Quét text -> danh sách từ nước ngoài chưa có trong kho, kèm cách đọc đề xuất.
    LƯU bản nháp kết quả (scan_draft.json) để F5 không làm mất. KHÔNG thêm vào kho."""
    data = request.get_json(force=True, silent=True) or {}
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"items": [], "total": 0})
    items = loanwords_dynamic.scan_new_words(text)
    _save_scan_draft(items)
    return jsonify({"items": items, "total": len(items)})


@app.route("/api/scan", methods=["GET"])
def api_scan_get():
    """Lấy bản nháp quét gần nhất (nếu có) — để khôi phục lại sau F5."""
    items = _load_scan_draft()
    return jsonify({"items": items, "total": len(items)})


@app.route("/api/scan/draft", methods=["POST"])
def api_scan_save_draft():
    """Tự lưu bản nháp khi LO sửa cách đọc trong bảng quét (body: {items:[...]})."""
    data = request.get_json(force=True, silent=True) or {}
    items = data.get("items") or []
    clean = []
    for it in items:
        if not isinstance(it, dict):
            continue
        word = (it.get("word") or "").strip()
        key = (it.get("key") or word.lower()).strip().lower()
        reading = (it.get("reading") or "").strip()
        if key and reading:
            clean.append({"word": word or key, "key": key, "reading": reading})
    _save_scan_draft(clean)
    return jsonify({"ok": True, "total": len(clean)})


@app.route("/api/scan/apply", methods=["POST"])
def api_scan_apply():
    """Thêm các từ đã duyệt (kèm cách đọc) vào kho động, đồng thời xóa chúng
    khỏi bản nháp quét. body: {items:[{word, reading}]}."""
    data = request.get_json(force=True, silent=True) or {}
    items = data.get("items") or []
    pairs = [(it.get("word", ""), it.get("reading", "")) for it in items]
    added = loanwords_dynamic.add_words(pairs)
    # xóa các từ đã thêm khỏi bản nháp
    added_keys = {p[0].strip().lower() for p in pairs if p[0] and p[1]}
    rest = [it for it in _load_scan_draft() if it.get("key") not in added_keys]
    if rest:
        _save_scan_draft(rest)
    else:
        _clear_scan_draft()
    return jsonify({
        "ok": True, "added": added,
        "total": loanwords_dynamic.count(),
        "draft_total": len(rest),
    })


@app.route("/api/scan", methods=["DELETE"])
def api_scan_clear():
    """Xóa bản nháp quét (khi LO bấm 'Xóa kết quả quét')."""
    _clear_scan_draft()
    return jsonify({"ok": True})


def extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    """Trích xuất nội dung văn bản từ các file: .txt, .md, .docx, .pdf, .srt, .json, v.v."""
    ext = Path(filename).suffix.lower()

    # 1. File Word (.docx)
    if ext == ".docx":
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " \t ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n".join(paragraphs).strip()
        except Exception:
            import zipfile
            import xml.etree.ElementTree as ET
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                xml_content = z.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                text_list = []
                for node in tree.iter():
                    if node.tag.endswith("t") and node.text:
                        text_list.append(node.text)
                    elif node.tag.endswith("p"):
                        text_list.append("\n")
                return "".join(text_list).strip()

    # 2. File PDF (.pdf)
    elif ext == ".pdf":
        try:
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            except ImportError:
                import PyPDF2 as pypdf
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages_text.append(t)
            return "\n\n".join(pages_text).strip()
        except Exception as e:
            raise RuntimeError(f"Không thể đọc file PDF: {str(e)}")

    # 3. File văn bản thuần: .txt, .md, .srt, .vtt, .json, .csv, .log, .html, v.v.
    for enc in ("utf-8", "utf-8-sig", "utf-16", "utf-16-le", "utf-16-be", "cp1258", "windows-1258", "latin-1"):
        try:
            return file_bytes.decode(enc).strip()
        except Exception:
            continue
    return file_bytes.decode("utf-8", errors="ignore").strip()


@app.route("/api/upload-file", methods=["POST"])
def api_upload_file():
    """Tải tệp lên và trích xuất nội dung văn bản (txt, md, docx, pdf, srt, v.v.)"""
    if "file" not in request.files:
        return jsonify({"error": "Không tìm thấy file tải lên"}), 400
    file = request.files["file"]
    if not file or not file.filename:
        return jsonify({"error": "File không hợp lệ"}), 400
    try:
        content = file.read()
        text = extract_text_from_file(file.filename, content)
        if not text:
            return jsonify({"error": "File trống hoặc không tìm thấy nội dung văn bản"}), 400
        return jsonify({
            "ok": True,
            "filename": file.filename,
            "text": text,
            "char_count": len(text),
        })
    except Exception as e:
        return jsonify({"error": f"Lỗi đọc file: {str(e)}"}), 500


@app.route("/san-pham/<path:name>")
def san_pham(name):
    return send_from_directory(str(PRODUCT_DIR), name)


@app.route("/uploads/<path:name>")
def uploads(name):
    return send_from_directory(str(PRODUCT_DIR), name)


@app.route("/api/history")
def api_history():
    items = []
    for f in sorted(PRODUCT_DIR.glob("*.mp3"), key=lambda p: p.stat().st_mtime, reverse=True):
        st = f.stat()
        dur = get_audio_duration(f)
        items.append({
            "name": f.name,
            "size_mb": round(st.st_size / 1024 / 1024, 2),
            "mtime": datetime.fromtimestamp(st.st_mtime).strftime("%d/%m/%Y %H:%M"),
            "url": f"/san-pham/{f.name}",
            "duration": round(dur, 2),
        })
    return jsonify(items)


@app.route("/api/history/<path:name>", methods=["DELETE"])
def api_history_delete(name):
    # chặn path traversal
    p = (PRODUCT_DIR / name).resolve()
    if not str(p).startswith(str(PRODUCT_DIR.resolve())):
        return jsonify({"error": "Không hợp lệ"}), 400
    if p.exists() and p.suffix == ".mp3":
        p.unlink()
        return jsonify({"ok": True})
    return jsonify({"error": "Không tìm thấy"}), 404


@app.route("/api/merge", methods=["POST"])
def api_merge():
    """Ghép nhiều file MP3 lại thành 1 file. body: {files: ["name1.mp3", ...], name?: "output.mp3"}"""
    data = request.get_json(force=True, silent=True) or {}
    files = data.get("files") or []
    if len(files) < 2:
        return jsonify({"error": "Cần ít nhất 2 file để ghép"}), 400

    ff = tts.FFMPEG
    if not ff:
        return jsonify({"error": "Không tìm thấy ffmpeg"}), 500

    # Kiểm tra tất cả file tồn tại
    paths = []
    for name in files:
        p = (PRODUCT_DIR / name).resolve()
        if not str(p).startswith(str(PRODUCT_DIR.resolve())):
            return jsonify({"error": f"File không hợp lệ: {name}"}), 400
        if not p.exists():
            return jsonify({"error": f"Không tìm thấy file: {name}"}), 404
        paths.append(p)

    # Tạo file danh sách concat
    list_txt = PRODUCT_DIR / f"_concat_{uuid.uuid4().hex[:8]}.txt"
    list_txt.write_text("\n".join(f"file '{p}'" for p in paths), encoding="utf-8")

    out_name = (data.get("name") or "").strip()
    if not out_name:
        out_name = f"Ghep_{datetime.now():%Y%m%d_%H%M%S}.mp3"
    if not out_name.lower().endswith(".mp3"):
        out_name += ".mp3"
    out_name = re.sub(r'[\\/*?:"<>|]', "", out_name)

    out_path = PRODUCT_DIR / out_name
    try:
        r = subprocess.run(
            [ff, "-y", "-f", "concat", "-safe", "0", "-i", str(list_txt),
             "-c", "copy", str(out_path)],
            capture_output=True, timeout=600,
        )
        if r.returncode != 0 or not out_path.exists():
            err = r.stderr.decode("utf-8", errors="ignore")[-500:]
            return jsonify({"error": f"ffmpeg lỗi: {err}"}), 500

        size_mb = round(out_path.stat().st_size / 1024 / 1024, 2)
        return jsonify({
            "ok": True,
            "name": out_name,
            "url": f"/san-pham/{out_name}",
            "size_mb": size_mb,
            "merged_count": len(paths),
        })
    finally:
        list_txt.unlink(missing_ok=True)


@app.route("/api/history/<path:name>/rename", methods=["POST"])
def api_history_rename(name):
    """Đổi tên file sản phẩm audio."""
    data = request.get_json(force=True, silent=True) or {}
    new_name = (data.get("new_name") or "").strip()
    if not new_name:
        return jsonify({"error": "Vui lòng nhập tên mới"}), 400

    clean_name = re.sub(r'[\\/*?:"<>|]', "", new_name).strip()
    if not clean_name.lower().endswith(".mp3"):
        clean_name += ".mp3"

    if not clean_name or clean_name == ".mp3":
        return jsonify({"error": "Tên file không hợp lệ"}), 400

    old_p = (PRODUCT_DIR / name).resolve()
    new_p = (PRODUCT_DIR / clean_name).resolve()

    if not str(old_p).startswith(str(PRODUCT_DIR.resolve())) or not str(new_p).startswith(str(PRODUCT_DIR.resolve())):
        return jsonify({"error": "Đường dẫn không hợp lệ"}), 400
    if not old_p.exists():
        return jsonify({"error": "Không tìm thấy file"}), 404
    if new_p.exists() and new_p != old_p:
        return jsonify({"error": "Tên file này đã tồn tại, vui lòng chọn tên khác"}), 409

    old_p.rename(new_p)
    return jsonify({"ok": True, "old_name": name, "new_name": clean_name, "url": f"/san-pham/{clean_name}"})


# ==================== VIDEO ====================

@app.route("/api/video/upload-image", methods=["POST"])
def api_video_upload_image():
    """Upload ảnh nền cho video. Trả về URL ảnh."""
    if "file" not in request.files:
        return jsonify({"error": "Không có file"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "File rỗng"}), 400
    ext = Path(f.filename).suffix.lower() or ".jpg"
    if ext not in (".jpg", ".jpeg", ".png", ".webp", ".bmp"):
        return jsonify({"error": "Chỉ hỗ trợ JPG, PNG, WEBP, BMP"}), 400
    out_name = f"img_{uuid.uuid4().hex[:8]}{ext}"
    out_path = UPLOAD_IMG_DIR / out_name
    f.save(str(out_path))
    return jsonify({"ok": True, "url": f"/upload-images/{out_name}", "name": out_name})


@app.route("/upload-images/<path:name>")
def serve_upload_image(name):
    return send_from_directory(str(UPLOAD_IMG_DIR), name)


@app.route("/api/video/preview-frame", methods=["POST"])
def api_video_preview_frame():
    """Tạo preview frame: ảnh + text overlay, trả về PNG base64."""
    data = request.get_json(force=True, silent=True) or {}
    img_url = data.get("image_url", "")
    text = data.get("text", "")
    text_color = data.get("text_color", "#FFFFFF")
    text_size = data.get("text_size", 48)
    text_y = data.get("text_y", 50)  # % từ trên xuống
    brightness = data.get("brightness", 100)  # %

    img_path = UPLOAD_IMG_DIR / Path(img_url).name
    if not img_path.exists():
        return jsonify({"error": "Ảnh không tồn tại"}), 404

    import base64
    # Tạo preview bằng ffmpeg
    filters = []
    if brightness != 100:
        filters.append(f"eq=brightness={(brightness - 100) / 100}")

    drawtext = f"drawtext=text='{text}':fontcolor={text_color}:fontsize={text_size}:x=(w-text_w)/2:y=h*{text_y}/100-text_h/2"
    if filters:
        filters.append(drawtext)
        vf = ",".join(filters)
    else:
        vf = drawtext

    cmd = [
        FFPROBE, "-v", "error", "-show_entries", "stream=width,height",
        "-of", "default=noprint_wrappers=1:nokey=1", str(img_path),
    ]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=5)
        wh = r.stdout.strip().split("\n")
        w, h = int(wh[0]), int(wh[1])
    except Exception:
        w, h = 1920, 1080

    tmp_out = UPLOAD_IMG_DIR / f"preview_{uuid.uuid4().hex[:8]}.png"
    cmd = [
        tts.FFMPEG, "-y", "-i", str(img_path),
        "-vf", vf, "-frames:v", "1", str(tmp_out),
    ]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
        if r.returncode != 0:
            return jsonify({"error": f"ffmpeg: {r.stderr[:300]}"}), 500
        with open(tmp_out, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        return jsonify({"ok": True, "image": f"data:image/png;base64,{b64}", "width": w, "height": h})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        tmp_out.unlink(missing_ok=True)


@app.route("/api/video/create", methods=["POST"])
def api_video_create():
    """Tạo video từ ảnh + audio + text overlay.
    body: {image_url, audio_url, text, text_color, text_size, text_y, brightness, output_name}"""
    data = request.get_json(force=True, silent=True) or {}
    img_url = data.get("image_url", "")
    audio_url = data.get("audio_url", "")
    text = data.get("text", "")
    text_color = data.get("text_color", "#FFFFFF")
    text_size = data.get("text_size", 48)
    text_y = data.get("text_y", 50)
    brightness = data.get("brightness", 100)
    output_name = data.get("output_name", "").strip()

    if not img_url or not audio_url:
        return jsonify({"error": "Cần ảnh + audio"}), 400

    img_path = (UPLOAD_IMG_DIR / Path(img_url).name).resolve()
    audio_path = (BASE_DIR / audio_url.lstrip("/")).resolve()
    if not img_path.exists():
        return jsonify({"error": "Ảnh không tồn tại"}), 404
    if not audio_path.exists():
        return jsonify({"error": "Audio không tồn tại"}), 404

    if not output_name:
        output_name = f"video_{datetime.now().strftime('%Y%m%d_%H%M%S')}.mp4"
    if not output_name.endswith(".mp4"):
        output_name += ".mp4"

    # Lấy duration audio
    dur = get_audio_duration(audio_path)
    if dur <= 0:
        return jsonify({"error": "Audio rỗng hoặc không đọc được"}), 400

    # Build ffmpeg filter
    filters = []
    if brightness != 100:
        filters.append(f"eq=brightness={(brightness - 100) / 100}")

    if text:
        safe_text = text.replace("'", "'\\''").replace(":", "\\:")
        filters.append(
            f"drawtext=text='{safe_text}':fontcolor={text_color}:fontsize={text_size}"
            f":x=(w-text_w)/2:y=h*{text_y}/100-text_h/2"
        )

    vf_parts = [f"scale=1920:1080:force_original_aspect_ratio=decrease,pad=1920:1080:(ow-iw)/2:(oh-ih)/2"]
    vf_parts.extend(filters)

    cmd = [
        tts.FFMPEG, "-y",
        "-loop", "1", "-i", str(img_path),
        "-i", str(audio_path),
        "-vf", ",".join(vf_parts),
        "-c:v", "libx264", "-preset", "fast", "-crf", "23",
        "-c:a", "aac", "-b:a", "128k",
        "-shortest", "-t", str(dur),
        "-pix_fmt", "yuv420p",
        str(VIDEO_DIR / output_name),
    ]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        if r.returncode != 0:
            return jsonify({"error": f"ffmpeg lỗi: {r.stderr[-500:]}"}), 500
    except subprocess.TimeoutExpired:
        return jsonify({"error": "ffmpeg timeout (>5 phút)"}), 500

    out_path = VIDEO_DIR / output_name
    if not out_path.exists():
        return jsonify({"error": "Video không được tạo ra"}), 500

    size_mb = round(out_path.stat().st_size / 1024 / 1024, 2)
    return jsonify({
        "ok": True,
        "name": output_name,
        "url": f"/video-output/{output_name}",
        "size_mb": size_mb,
        "duration": round(dur, 2),
    })


@app.route("/video-output/<path:name>")
def serve_video(name):
    p = (VIDEO_DIR / name).resolve()
    if not str(p).startswith(str(VIDEO_DIR.resolve())):
        return "Not found", 404
    return send_from_directory(str(VIDEO_DIR), name)


@app.route("/api/video/list")
def api_video_list():
    items = []
    for f in sorted(VIDEO_DIR.glob("*.mp4"), key=lambda p: p.stat().st_mtime, reverse=True):
        st = f.stat()
        dur = get_audio_duration(f)
        items.append({
            "name": f.name,
            "size_mb": round(st.st_size / 1024 / 1024, 2),
            "mtime": datetime.fromtimestamp(st.st_mtime).strftime("%d/%m/%Y %H:%M"),
            "url": f"/video-output/{f.name}",
            "duration": round(dur, 2),
        })
    return jsonify(items)


@app.route("/api/video/<path:name>", methods=["DELETE"])
def api_video_delete(name):
    p = (VIDEO_DIR / name).resolve()
    if not str(p).startswith(str(VIDEO_DIR.resolve())):
        return jsonify({"error": "Không hợp lệ"}), 400
    if p.exists() and p.suffix == ".mp4":
        p.unlink()
        return jsonify({"ok": True})
    return jsonify({"error": "Không tìm thấy"}), 404


if __name__ == "__main__":
    import threading as _th
    import webbrowser as _wb

    def _open_browser():
        import time as _t
        _t.sleep(1.5)
        _wb.open("http://127.0.0.1:5000/")

    _th.Thread(target=_open_browser, daemon=True).start()
    app.run(host="127.0.0.1", port=5000, debug=False, threaded=True)